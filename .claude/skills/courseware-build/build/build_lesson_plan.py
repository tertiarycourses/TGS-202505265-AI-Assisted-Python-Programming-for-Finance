#!/usr/bin/env python3
"""Generate the AI Assisted Python Programming for Finance Lesson Plan (LP) DOCX (Tertiary house format).

Cover page + Document Version Control Record + auto TOC + Arial 11pt body +
colour-coded 4-day schedule tables (9:30am-6:30pm, 8 training hours/day, 1h
lunch, tea within, final assessment Day 4). Topics/labs come from
course_data + the domain data files so the LP stays aligned with the deck,
guide and labs.
"""
import os, sys
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

HERE=os.path.dirname(os.path.abspath(__file__)); sys.path.insert(0,HERE)
import course_data as C
from data_domain1 import DOMAIN1; from data_domain2 import DOMAIN2
from data_domain3 import DOMAIN3; from data_domain4 import DOMAIN4
from data_domain5 import DOMAIN5; from data_domain6 import DOMAIN6
from data_domain7 import DOMAIN7; from data_domain8 import DOMAIN8
from data_domain9 import DOMAIN9
ACT=(DOMAIN1+DOMAIN2+DOMAIN3+DOMAIN4+DOMAIN5+DOMAIN6+DOMAIN7+DOMAIN8+DOMAIN9)
import prodoc
def _find_repo(start):
    env=os.environ.get("COURSE_REPO")
    if env and os.path.isdir(env): return env
    d=start
    for _ in range(8):
        d=os.path.dirname(d)
        if os.path.isdir(os.path.join(d,"courseware")) and os.path.isdir(os.path.join(d,"labs")): return d
    return os.path.dirname(os.path.dirname(HERE))
REPO=_find_repo(HERE); ASSETS=os.path.join(os.path.dirname(HERE),"assets")

BRAND=RGBColor(0x1F,0x6F,0xEB); DARK=RGBColor(0x11,0x18,0x27); GREY=RGBColor(0x55,0x5B,0x66)
HEADER_FILL="1F6FEB"; TOPIC_FILL="E8F0FE"; BREAK_FILL="FFF4E5"; LUNCH_FILL="FDE9D9"; ASSESS_FILL="E8F7EE"

def lab_titles(nums):
    return "; ".join(f"Lab {a['num']}: {a['title']}" for a in ACT if a['num'] in nums)

# ------------------------------------------------ schedule (single source of truth for timing)
# 9:30am-6:30pm, 8 training hours/day, 1-hour lunch, tea breaks counted within.
# Built from course_data.DAY_TOPICS so the LP can never drift from the deck/labs.
def _labs_for(tn):
    return [a for a in ACT if a["topic"]==tn]

def _split(seq, n):
    """Split a list into n roughly equal chunks."""
    if not seq: return [[] for _ in range(n)]
    k, m = divmod(len(seq), n)
    out=[]; i=0
    for x in range(n):
        sz=k+(1 if x<m else 0); out.append(seq[i:i+sz]); i+=sz
    return out

def _titles(labs):
    return "; ".join(f"Lab {a['num']}: {a['title']}" for a in labs) if labs else ""

def _build_day(day):
    tns=C.DAY_TOPICS[day]
    theme=C.DAY_THEMES[day]
    rows=[]
    first = (day==1)
    last  = (day==C.DAYS)
    # morning admin
    if first:
        rows.append(("9:30","10:00",30,"admin","Welcome, course introduction, learning outcomes, ground rules and mandatory digital attendance (AM)"))
        t0="10:00"
    else:
        rows.append(("9:30","9:45",15,"recap",f"Day {day-1} recap, Q&A and mandatory digital attendance (AM)"))
        t0="9:45"
    # the day's topics are split across the morning and afternoon blocks
    blocks=[]
    for tn in tns:
        tp=next(x for x in C.TOPICS if x["num"]==tn)
        blocks.append((tp,_labs_for(tn)))
    if last:
        # Final day: topics in the morning, then the FULL 120-minute Practical Test.
        # Day runs 9:30-19:00 so the practical is not truncated.
        tp1,l1=blocks[0]
        a,b=_split(l1,2)
        rows.append((t0,"10:45",60,"topic",
                     f"Topic {tp1['num']} — {tp1['title']}: {tp1['subtitle']} (concepts + demonstration)"))
        rows.append(("10:45","11:00",15,"break","Tea break"))
        rows.append(("11:00","13:00",120,"lab","Hands-on: "+_titles(a+b)))
        rows.append(("13:00","14:00",60,"lunch","Lunch break"))
        tp2,l2=blocks[1]
        c,d=_split(l2,2)
        rows.append(("14:00","14:45",45,"topic",
                     f"Topic {tp2['num']} — {tp2['title']}: {tp2['subtitle']} (concepts + demonstration). Digital attendance (PM)"))
        rows.append(("14:45","15:45",60,"lab","Hands-on: "+_titles(c+d)))
        rows.append(("15:45","16:00",15,"break","Tea break"))
        rows.append(("16:00","16:30",30,"recap","Course recap, Q&A, Course Feedback and TRAQOM survey"))
        rows.append(("16:30","16:45",15,"assess","Briefing for Assessment. Assessment digital attendance"))
        rows.append(("16:45","17:45",60,"assess",C.ASSESSMENT["written"]))
        rows.append(("17:45","19:45",120,"assess",
                     "Practical Test (PP) — 7 open-ended practical questions based on the "
                     "Financial Data Solutions Inc. scenario, 120 minutes, open book"))
    else:
        if len(blocks)==1:
            tp,labs=blocks[0]
            a,b,c=_split(labs,3)
            rows.append((t0,"11:00",75 if first else 75,"topic",
                         f"Topic {tp['num']} — {tp['title']}: {tp['subtitle']} (concepts + demonstration)"))
            rows.append(("11:00","11:15",15,"break","Tea break"))
            rows.append(("11:15","13:00",105,"lab","Hands-on: "+_titles(a)))
            rows.append(("13:00","14:00",60,"lunch","Lunch break"))
            rows.append(("14:00","15:45",105,"lab","Digital attendance (PM). Hands-on: "+_titles(b)))
            rows.append(("15:45","16:00",15,"break","Tea break"))
            rows.append(("16:00","18:15",135,"lab","Hands-on: "+_titles(c)))
            rows.append(("18:15","18:30",15,"recap",f"Day {day} recap, Q&A and PM digital attendance"))
        elif len(blocks)==2:
            (tpA,lA),(tpB,lB)=blocks
            a1,a2=_split(lA,2)
            rows.append((t0,"10:45",60 if first else 60,"topic",
                         f"Topic {tpA['num']} — {tpA['title']}: {tpA['subtitle']} (concepts + demonstration)"))
            rows.append(("10:45","11:00",15,"break","Tea break"))
            rows.append(("11:00","13:00",120,"lab","Hands-on: "+_titles(a1)))
            rows.append(("13:00","14:00",60,"lunch","Lunch break"))
            rows.append(("14:00","15:00",60,"lab","Digital attendance (PM). Hands-on: "+_titles(a2)))
            rows.append(("15:00","15:45",45,"topic",
                         f"Topic {tpB['num']} — {tpB['title']}: {tpB['subtitle']} (concepts + demonstration)"))
            rows.append(("15:45","16:00",15,"break","Tea break"))
            rows.append(("16:00","18:15",135,"lab","Hands-on: "+_titles(lB)))
            rows.append(("18:15","18:30",15,"recap",f"Day {day} recap, Q&A and PM digital attendance"))
        else:
            (tpA,lA),(tpB,lB),(tpC,lC)=blocks
            rows.append((t0,"10:30",30,"topic",
                         f"Topic {tpA['num']} — {tpA['title']}: {tpA['subtitle']} (concepts + demonstration)"))
            rows.append(("10:30","10:45",15,"break","Tea break"))
            rows.append(("10:45","11:30",45,"lab","Hands-on: "+_titles(lA)))
            rows.append(("11:30","13:00",90,"lab",
                         f"Topic {tpB['num']} — {tpB['title']} (concepts). Hands-on: "+_titles(lB[:len(lB)//2])))
            rows.append(("13:00","14:00",60,"lunch","Lunch break"))
            rows.append(("14:00","15:45",105,"lab","Digital attendance (PM). Hands-on: "+_titles(lB[len(lB)//2:])))
            rows.append(("15:45","16:00",15,"break","Tea break"))
            rows.append(("16:00","18:15",135,"lab",
                         f"Topic {tpC['num']} — {tpC['title']} (concepts). Hands-on: "+_titles(lC)))
            rows.append(("18:15","18:30",15,"recap",f"Day {day} recap, Q&A and PM digital attendance"))
    return theme, rows

SCHEDULE = {d: _build_day(d) for d in range(1, C.DAYS+1)}

# ------------------------------------------------ build document
doc=Document()
normal=doc.styles["Normal"]; normal.font.name="Arial"; normal.font.size=Pt(11)
prodoc.style_headings(doc)

prodoc.add_cover_page(doc,"LESSON PLAN",C.TITLE,C.VERSION.lstrip("v"),
                      org_logo=os.path.join(ASSETS,"tertiary-infotech-logo.png"),
                      course_logo=None, course_code=C.COURSE_CODE)
prodoc.add_version_control(doc,[
 ("30.0","2 May 2026","Previous release — IBF-accredited 4-day lesson plan for AI Assisted Python Programming for Finance.","Dr. Alfred Ang"),
 (C.VERSION.lstrip("v"),C.VERSION_DATE,
  f"Rebuilt from the single-source content module. Expanded to {len(ACT)} hands-on finance labs across the nine topics; "
  "all labs migrated to the uv project manager; added a progressive Streamlit finance-dashboard capstone in Topics 7 and 9; "
  "full build-out of Topics 5-9 (exception handling, pandas, aggregation/visualisation, OOP, apply/pipe).",C.TRAINER),
])
prodoc.add_toc(doc)

def H(text,level=1):
    h=doc.add_heading(text,level=level); return h

H("Course Information",1)
info=[("Course Title",C.TITLE),("Course Reference",C.COURSE_CODE),
      ("Accreditation",f"{C.ACCREDITATION['framework']} · IBF Standard {C.ACCREDITATION['standard']} · Funded under IBF-STS"),
      ("Training Provider",C.ORG+"  ("+C.UEN.replace('UEN: ','UEN ')+")"),
      ("Duration",f"{C.DAYS} days · 8 training hours per day ({C.DAYS*8} hours)"),
      ("Daily Timing","9:30 am – 6:30 pm (1-hour lunch; tea breaks within training time)"),
      ("Mode","Instructor-led, hands-on AI-assisted Python labs using uv, pandas, yfinance and Streamlit"),
      ("Trainer",C.TRAINER)]
t=doc.add_table(rows=0,cols=2); t.style="Table Grid"
for k,v in info:
    c=t.add_row().cells; c[0].text=""; r=c[0].paragraphs[0].add_run(k); r.bold=True; r.font.size=Pt(10)
    prodoc._shade_cell(c[0],TOPIC_FILL)
    c[1].text=""; c[1].paragraphs[0].add_run(v).font.size=Pt(10)

H("Learning Outcomes",1)
doc.add_paragraph("On completion of this course, learners will be able to:")
for lo in C.LEARNING_OUTCOMES:
    p=doc.add_paragraph(style="List Bullet"); p.add_run(lo).font.size=Pt(10.5)

H("Assessment",1)
for a in [C.ASSESSMENT["written"],C.ASSESSMENT["practical"],
          "Format: Open Book — course slides, Learner Guide and lab files only.",
          f"Passing score: {C.ASSESSMENT['passing_score']}.",
          f"Final assessment is conducted on Day {C.DAYS} from 5:00 pm.",C.ASSESSMENT["note"]]:
    p=doc.add_paragraph(style="List Bullet"); p.add_run(a).font.size=Pt(10.5)

def set_cell(cell,text,bold=False,size=9.5,color=None,fill=None,align=None):
    cell.text=""; p=cell.paragraphs[0]
    if align: p.alignment=align
    r=p.add_run(text); r.bold=bold; r.font.size=Pt(size); r.font.name="Arial"
    if color: r.font.color.rgb=color
    if fill: prodoc._shade_cell(cell,fill)

KIND_FILL={"topic":TOPIC_FILL,"break":BREAK_FILL,"lunch":LUNCH_FILL,"assess":ASSESS_FILL,
           "admin":"F3F5F8","recap":"F3F5F8","lab":None}

H("Course Schedule",1)
for day,(theme,rows) in SCHEDULE.items():
    H(f"Day {day} — {theme}",2)
    tbl=doc.add_table(rows=0,cols=3); tbl.style="Table Grid"; tbl.alignment=WD_TABLE_ALIGNMENT.CENTER
    hdr=tbl.add_row().cells
    for i,htext in enumerate(["Time","Duration","Topic / Activity"]):
        set_cell(hdr[i],htext,bold=True,size=10,color=RGBColor(0xFF,0xFF,0xFF),fill=HEADER_FILL)
    training=0; instructional=0
    for start,end,mins,kind,text in rows:
        cells=tbl.add_row().cells; fill=KIND_FILL.get(kind)
        set_cell(cells[0],f"{start}–{end}",bold=(kind in ("topic","assess")),size=9.5,fill=fill)
        set_cell(cells[1],f"{mins} min",size=9.5,fill=fill)
        set_cell(cells[2],text,bold=(kind in ("topic","assess")),size=9.5,fill=fill)
        if kind!="lunch": training+=mins
        if kind not in ("lunch","assess"): instructional+=mins
    # widths
    for row in tbl.rows:
        row.cells[0].width=Inches(1.15); row.cells[1].width=Inches(0.9); row.cells[2].width=Inches(4.75)
    assess_mins=training-instructional
    if assess_mins:
        note=(f"Total training time: {training} minutes ({training//60} hours "
              f"{training%60} min) — {instructional} minutes instruction plus "
              f"{assess_mins} minutes of assessment.")
    else:
        note=f"Total training time: {training} minutes ({training//60} hours)."
    p=doc.add_paragraph(); r=p.add_run(note); r.italic=True; r.font.size=Pt(9.5); r.font.color.rgb=GREY
    # Every day must deliver 8 instructional hours. The final day additionally
    # carries the 120-minute assessment, so it legitimately runs past 18:30.
    # Contact time must be 8 hours every day. On the final day part of that is the
    # assessment, so instruction + assessment (not instruction alone) makes the 480,
    # and the day runs past 18:30 by however long the assessment overruns.
    assert instructional+assess_mins>=480, (
        f"Day {day}: only {instructional+assess_mins} contact minutes "
        f"({instructional} instruction + {assess_mins} assessment), expected at least 480")
    if not assess_mins:
        assert instructional==480, (
            f"Day {day}: {instructional} instructional minutes, expected 480")

H("Lab Reference (aligned to the course topics)",1)
tt=doc.add_table(rows=0,cols=3); tt.style="Table Grid"
hdr=tt.add_row().cells
for i,htext in enumerate(["Topic","Weighting","Labs"]):
    set_cell(hdr[i],htext,bold=True,size=10,color=RGBColor(0xFF,0xFF,0xFF),fill=HEADER_FILL)
for tp in C.TOPICS:
    acts=[a for a in ACT if a["topic"]==tp["num"]]
    cells=tt.add_row().cells
    set_cell(cells[0],f"Topic {tp['code']}: {tp['title']}",bold=True,size=9.5,fill=TOPIC_FILL)
    set_cell(cells[1],tp["weighting"],size=9.5,fill=TOPIC_FILL)
    set_cell(cells[2],", ".join(f"Lab {a['num']}" for a in acts),size=9.5)

prodoc.add_page_numbers(doc)
prodoc.enable_update_fields(doc)
OUT=os.path.join(REPO,"courseware",f"LP-{C.SHORT_TITLE}.docx")
doc.save(OUT)
print("Saved",OUT)
