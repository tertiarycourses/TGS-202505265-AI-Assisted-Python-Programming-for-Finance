#!/usr/bin/env python3
"""Generate the AI Assisted Python Programming for Finance Learner Guide as BOTH a Markdown mirror (LG-*.md at repo
root) and a DOCX (courseware/LG-*.docx) from one source, so they never diverge.

House format: cover page, Document Version Control Record, auto TOC, Arial 11pt
body, one section per lab (Objective · Goal · What you'll build · Step-by-step
with commands · Test it), plus setup, exam-prep and glossary. All content is
driven by course_data + the domain data files, keeping the LG 100% aligned with
the slide deck, Lesson Plan and labs.
"""
import os, sys
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

HERE=os.path.dirname(os.path.abspath(__file__)); sys.path.insert(0,HERE)
import course_data as C
from data_domain1 import DOMAIN1; from data_domain2 import DOMAIN2
from data_domain3 import DOMAIN3; from data_domain4 import DOMAIN4
from data_domain5 import DOMAIN5; from data_domain6 import DOMAIN6
from data_domain7 import DOMAIN7; from data_domain8 import DOMAIN8
from data_domain9 import DOMAIN9
ACT=(DOMAIN1+DOMAIN2+DOMAIN3+DOMAIN4+DOMAIN5+DOMAIN6+DOMAIN7+DOMAIN8+DOMAIN9)
import prodoc
def _find_repo(start):
    env=os.environ.get("COURSE_REPO")
    if env and os.path.isdir(env): return env
    d=start
    for _ in range(8):
        d=os.path.dirname(d)
        if os.path.isdir(os.path.join(d,"courseware")) and os.path.isdir(os.path.join(d,"labs")): return d
    return os.path.dirname(os.path.dirname(HERE))
REPO=_find_repo(HERE); ASSETS=os.path.join(os.path.dirname(HERE),"assets")

# ---------------- block DSL (single content stream → MD + DOCX) ----------------
B=[]
def h1(t): B.append(("h1",t))
def h2(t): B.append(("h2",t))
def h3(t): B.append(("h3",t))
def p(t):  B.append(("p",t))
def bullets(xs): B.append(("bullets",xs))
def steps(xs): B.append(("steps",xs))
def code(t): B.append(("code",t))
def note(t): B.append(("note",t))
def rule(): B.append(("rule",))

# ---------------- content ----------------
h1("Introduction")
p(f"This Learner Guide accompanies the course {C.TITLE} ({C.COURSE_CODE}), conducted by {C.ORG}. "
  f"The course is accredited under the {C.ACCREDITATION['framework']}, IBF Standard {C.ACCREDITATION['standard']}, "
  f"and is eligible for funding under the IBF Standards Training Scheme (IBF-STS). "
  f"This guide provides step-by-step instructions for all {len(ACT)} hands-on labs, organised by the "
  f"{len(C.TOPICS)} course topics.")
p("Every lab is grounded in a real financial-services task — CPF and loan calculations, SGX market data, "
  "financial ratios such as ROE, ROA and net profit margin, credit-risk classification, algorithmic-trading "
  "backtests and portfolio analytics. You will use AI assistants (Google Colab's Gemini, GitHub Copilot in "
  "VS Code, or Cursor) to draft code quickly, and then review, correct and test every line before relying on it.")
p("Use this guide alongside the course slides and the lab files in the labs/ folder. Each lab is also provided "
  "as a standalone Python file, and all labs are combined into a single Jupyter notebook you can open in "
  "Google Colab.")

h1("Course Learning Outcomes")
bullets(C.LEARNING_OUTCOMES)

h1("About IBF-STS Accreditation and Funding")
p(C.ACCREDITATION["blurb"])
p(f"Find out more at {C.ACCREDITATION['website']}.")

h1("Before You Start — Environment Setup")
h3("What you need")
bullets([
 "A computer with internet access, or a Google account for Google Colab (browser-based, nothing to install).",
 "Python 3.12 or later for local work, installed together with uv.",
 "An AI-assisted editor: Google Colab (built-in Gemini), Visual Studio Code with GitHub Copilot, Cursor, or AntiGravity.",
 "An internet connection for the labs that download live market data through the yfinance package.",
])
h3("Install uv — the project manager used in every lab")
p("This course uses uv, a single fast tool that replaces pip, virtualenv and pyenv. uv creates an isolated "
  "environment per lab and records exact dependency versions, so your results are reproducible — an audit "
  "requirement in financial services. Install it once, then verify:")
code("# macOS / Linux\ncurl -LsSf https://astral.sh/uv/install.sh | sh\n\n# Windows (PowerShell)\npowershell -ExecutionPolicy ByPass -c \"irm https://astral.sh/uv/install.ps1 | iex\"\n\n# verify the installation\nuv --version")
h3("The three uv commands you will use in every lab")
code("uv init lab-01-hello-finance      # create a new project folder + pyproject.toml\ncd lab-01-hello-finance\nuv add pandas yfinance matplotlib # install and record dependencies\nuv run python main.py             # run inside the project environment")
h3("Working in Google Colab instead")
p("If you prefer Colab, open a new notebook and install packages in a cell with a leading exclamation mark "
  "(for example !pip install yfinance). The lab logic is identical; only the environment setup differs. "
  "The combined notebook for this course is provided in the Activities folder.")
h3("Conventions used in every lab")
bullets([
 "Commands prefixed with $ are typed into your terminal; the $ itself is not part of the command.",
 "Each lab is a self-contained uv project in its own folder, named lab-NN-<short-title>.",
 "Where a lab depends on a file produced by an earlier lab, the step says so explicitly and copies it forward.",
 "AI-prompting steps show the exact prompt in quotation marks — type it, then review the generated code against the review step that follows.",
 "Financial figures use Singapore dollars unless a lab states otherwise; market data is live, so exact values will differ from the examples.",
])
h3("A note on AI-generated code")
p("An AI assistant will produce plausible-looking code very quickly. In financial services, plausible is not "
  "the same as correct: a wrong sign, an off-by-one on a rolling window, or a look-ahead bias in a backtest "
  "produces numbers that look reasonable and are wrong. Every lab therefore pairs an AI-prompting step with a "
  "review step. Never submit or deploy generated code you have not read, run and tested.")

# ---------------- per-topic, per-lab ----------------
TOPICS_BY_NUM={t["num"]:t for t in C.TOPICS}
for t in C.TOPICS:
    h1(f"Topic {t['code']} — {t['title']}  ({t['weighting']})")
    p(t["subtitle"])
    h3("Key concepts")
    bullets(t["concepts"])
    for a in [x for x in ACT if x["topic"]==t["num"]]:
        h2(f"Lab {a['num']} — {a['title']}")
        p(f"Learning objective: {a['objective']}")
        p(f"Goal: {a['desc']}")
        h3("What you'll build")
        p(a["build"]+f"   (Tools: {a['services']}.)")
        h3("Step-by-step")
        st=[]
        for i,(instr,cmd) in enumerate(a["steps"],1):
            st.append((instr,cmd))
        steps(st)
        h3("Test it")
        p(a["test"])
        note(f"The runnable source for this lab is labs/lab-{a['num']:02d}-*.py, and it is also included "
             f"as a section of the combined notebook labs/{C.SHORT_TITLE}-All-Labs.ipynb.")
        rule()

h1("Financial Formula Reference")
p("These are the formulas used across the labs and the assessment. Know what each measures and when the "
  "denominator can be zero — that is where both bugs and exception handling matter.")
gl_f=[
 ("Return on Equity (ROE)","Net Income / Shareholders' Equity x 100%. Measures profit generated per dollar of equity. Zero or negative equity makes this undefined or misleading."),
 ("Return on Assets (ROA)","Net Income / Total Assets x 100%. Measures how efficiently assets generate profit."),
 ("Net Profit Margin","Net Income / Revenue x 100%. The share of each revenue dollar that becomes profit. Guard against zero revenue."),
 ("Gross Profit Margin","(Revenue - COGS) / Revenue x 100%. Profitability before operating expenses."),
 ("Price-to-Earnings (PE) Ratio","Share Price / Earnings Per Share. A valuation multiple; negative earnings make it meaningless."),
 ("Debt-to-Income Ratio","Total Debt Repayments / Gross Income. A core credit-risk measure in loan assessment."),
 ("Compound Interest","A = P x (1 + r/n)^(n x t), where P is principal, r the annual rate, n the compounding frequency and t the years."),
 ("Loan Monthly Instalment","P x [r(1+r)^n] / [(1+r)^n - 1], where r is the monthly rate and n the number of payments."),
 ("Future Value of Regular Savings","P x [((1 + r)^n - 1) / r], used in the retirement calculator lab."),
 ("Year-on-Year Growth","(Current - Previous) / Previous x 100%. Guard against a zero or missing previous value."),
 ("Simple Moving Average (SMA)","The mean closing price over the last N periods; the basis of the crossover strategy in Lab 42."),
 ("Daily Return","(Today's Close - Yesterday's Close) / Yesterday's Close. Chained returns give cumulative performance."),
]
B.append(("dl",gl_f))
rule()

h1("Assessment Preparation")
p(f"The final assessment is conducted on Day {C.DAYS} and is open book — you may use these slides, this "
  f"Learner Guide and your lab files. The passing score is {C.ASSESSMENT['passing_score']}.")
bullets([
 C.ASSESSMENT["written"],
 C.ASSESSMENT["practical"],
 "The practical questions are based on a financial-services scenario with revenue, stock price, monthly sales and growth-rate datasets supplied in the paper.",
 "Expect to break a business requirement into functions, calculate growth rates, forecast revenue, classify price movements, debug and optimise a broken function, and document a function properly.",
 "Practise writing a complete function from a plain-English requirement without an AI assistant — you must be able to read and correct code unaided.",
 "Revisit any lab whose 'Test it' check you could not pass from memory.",
 "Sharpen your readiness with the Tertiary Infotech practice exams at https://exams.tertiaryinfotech.com",
])
rule()

h1("Glossary")
gl=[
 ("uv","A fast Python package and project manager used in every lab; replaces pip, virtualenv and pyenv."),
 ("pyproject.toml","The project file uv creates to record your dependencies, making the environment reproducible."),
 ("Series / DataFrame","The two core pandas structures: a single labelled column, and a labelled table of columns."),
 ("yfinance","A Python package that downloads historical and current market data from Yahoo Finance."),
 ("Ticker","The exchange symbol identifying a listed instrument, e.g. D05.SI for DBS Group on the SGX."),
 ("Market Capitalisation","Share price multiplied by shares outstanding — the market value of a company's equity."),
 ("Comprehension","Compact Python syntax that builds a list, set or dict in one expression, e.g. [f(x) for x in xs]."),
 ("Lambda","A small anonymous function written inline, commonly passed to map() or filter()."),
 ("Exception","An event during execution that disrupts normal flow, handled with try/except so a pipeline survives bad data."),
 ("Traceback","The error report Python prints when an exception is unhandled; it names the file, line and error type."),
 ("groupby","A pandas operation that splits data into groups, applies an aggregation, and combines the results."),
 ("Pivot table","A cross-tabulation that summarises one metric across two categorical dimensions."),
 ("apply / pipe","pandas methods that run a function across rows or columns (apply) or chain whole-DataFrame transformations (pipe)."),
 ("Class / Object","A class is the blueprint; an object is one instance of it — for example one Stock object per holding."),
 ("Inheritance","Deriving a specialised class (Equity, Bond) from a general one (Instrument) without duplicating code."),
 ("Polymorphism","One method call behaving correctly across different object types, e.g. .value() on any instrument."),
 ("Streamlit","A Python framework that turns an analysis script into an interactive web application."),
 ("Backtest","Running a trading strategy over historical data to estimate how it would have performed."),
 ("Look-ahead bias","A backtesting error where the strategy uses information that was not available at the time of the trade."),
]
B.append(("dl",gl))

# ---------------- render Markdown ----------------
def _anchor(txt):
    return "".join(ch.lower() if ch.isalnum() else ("-" if ch in " -" else "") for ch in txt)

def render_md():
    out=[f"# {C.TITLE} — Learner Guide",""]
    out.append(f"**Course Code:** {C.COURSE_CODE}  |  **Conducted by:** {C.ORG} ({C.UEN.replace('UEN: ','UEN ')})  |  **Version {C.VERSION} · {C.VERSION_DATE}**")
    out.append("")
    # TOC (h1 + h2)
    out.append("## Contents"); out.append("")
    for kind,*rest in B:
        if kind=="h1": out.append(f"- [{rest[0]}](#{_anchor(rest[0])})")
        elif kind=="h2": out.append(f"  - [{rest[0]}](#{_anchor(rest[0])})")
    out.append("")
    for kind,*rest in B:
        if kind=="h1": out+=["",f"## {rest[0]}",""]
        elif kind=="h2": out+=["",f"### {rest[0]}",""]
        elif kind=="h3": out+=[f"**{rest[0]}**",""]
        elif kind=="p": out+=[rest[0],""]
        elif kind=="bullets": out+=[f"- {x}" for x in rest[0]]+[""]
        elif kind=="steps":
            for i,(instr,cmd) in enumerate(rest[0],1):
                out.append(f"{i}. {instr}")
                if cmd: out+=["",f"   ```bash",f"   {cmd}","   ```",""]
            out.append("")
        elif kind=="code": out+=["```bash",rest[0],"```",""]
        elif kind=="note": out+=[f"> **Note:** {rest[0]}",""]
        elif kind=="rule": out+=["---",""]
        elif kind=="dl":
            for term,defn in rest[0]: out.append(f"- **{term}** — {defn}")
            out.append("")
    return "\n".join(out)

MD_OUT=os.path.join(REPO,f"LG-{C.SHORT_TITLE}.md")
with open(MD_OUT,"w") as f: f.write(render_md())
print("Saved",MD_OUT)

# ---------------- render DOCX ----------------
BRAND=RGBColor(0x1F,0x6F,0xEB); DARK=RGBColor(0x11,0x18,0x27); GREY=RGBColor(0x55,0x5B,0x66)
INKCODE=RGBColor(0x0B,0x30,0x60)
doc=Document()
normal=doc.styles["Normal"]; normal.font.name="Arial"; normal.font.size=Pt(11)
prodoc.style_headings(doc)
prodoc.add_cover_page(doc,"LEARNER GUIDE",C.TITLE,C.VERSION.lstrip("v"),
                      org_logo=os.path.join(ASSETS,"tertiary-infotech-logo.png"),
                      course_logo=None, course_code=C.COURSE_CODE)
prodoc.add_version_control(doc,[
 ("30.0","2 May 2026","Previous release — IBF-accredited Learner Guide for AI Assisted Python Programming for Finance.","Dr. Alfred Ang"),
 (C.VERSION.lstrip("v"),C.VERSION_DATE,
  f"Rebuilt from the single-source content module. Expanded to {len(ACT)} hands-on finance labs across the nine topics; "
  "all labs migrated to the uv project manager; added a progressive Streamlit finance-dashboard capstone; "
  "full build-out of Topics 5-9; added a Financial Formula Reference and an expanded glossary.",C.TRAINER),
])
prodoc.add_toc(doc)

def code_para(text):
    for line in text.split("\n"):
        para=doc.add_paragraph(); prodoc._shade_para(para) if hasattr(prodoc,"_shade_para") else None
        r=para.add_run(line); r.font.name="Consolas"; r.font.size=Pt(9.5); r.font.color.rgb=INKCODE

for kind,*rest in B:
    if kind=="h1": doc.add_heading(rest[0],level=1)
    elif kind=="h2": doc.add_heading(rest[0],level=2)
    elif kind=="h3":
        para=doc.add_paragraph(); r=para.add_run(rest[0]); r.bold=True; r.font.size=Pt(11); r.font.color.rgb=BRAND
    elif kind=="p": doc.add_paragraph(rest[0])
    elif kind=="bullets":
        for x in rest[0]: doc.add_paragraph(x,style="List Bullet")
    elif kind=="steps":
        for i,(instr,cmd) in enumerate(rest[0],1):
            para=doc.add_paragraph(style="List Number"); para.add_run(instr)
            if cmd: code_para(cmd)
    elif kind=="code": code_para(rest[0])
    elif kind=="note":
        para=doc.add_paragraph(); r=para.add_run("Note: "); r.bold=True; r.font.color.rgb=BRAND
        para.add_run(rest[0]).font.size=Pt(10)
    elif kind=="rule": doc.add_paragraph("")
    elif kind=="dl":
        for term,defn in rest[0]:
            para=doc.add_paragraph(style="List Bullet")
            r=para.add_run(term+" — "); r.bold=True; para.add_run(defn)

prodoc.add_page_numbers(doc)
prodoc.enable_update_fields(doc)
DOCX_OUT=os.path.join(REPO,"courseware",f"LG-{C.SHORT_TITLE}.docx")
doc.save(DOCX_OUT)
print("Saved",DOCX_OUT)
